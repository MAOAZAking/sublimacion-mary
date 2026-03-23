import os
import json
from docx import Document
from docx.shared import Mm, Cm
from docx.enum.section import WD_ORIENT

# Cargar pedidos.json para saber que imagenes son las validas
def load_valid_images():
    valid_images = set()
    try:
        # Asumiendo que se ejecuta desde la raiz
        json_path = os.path.join("json", "pedidos.json")
        if not os.path.exists(json_path):
            return None
            
        with open(json_path, 'r', encoding='utf-8') as f:
            pedidos = json.load(f)
            
        for p in pedidos:
            # Extraer nombres de archivo de las URLs
            if "imagen_url" in p and p["imagen_url"]:
                valid_images.add(os.path.basename(p["imagen_url"]))
            
            if "imagenes" in p and p["imagenes"]:
                imgs = p["imagenes"]
                if isinstance(imgs, dict):
                    if "frontal" in imgs and imgs["frontal"]: valid_images.add(os.path.basename(imgs["frontal"]))
                    if "espaldar" in imgs and imgs["espaldar"]: valid_images.add(os.path.basename(imgs["espaldar"]))
                    if "lamina" in imgs and imgs["lamina"]: valid_images.add(os.path.basename(imgs["lamina"]))
                    
    except Exception as e:
        print(f"Advertencia: No se pudo cargar pedidos.json para filtrar imagenes: {e}")
        return None
    return valid_images

def create_word_for_folder(folder_path, valid_images_set):
    # Buscar imagenes validas (laminas)
    images = []
    for f in os.listdir(folder_path):
        lower_f = f.lower()
        # Criterio: Empieza con lamina_, no es preview y es una imagen
        if lower_f.startswith('lamina_') and not '_preview' in lower_f and (lower_f.endswith('.png') or lower_f.endswith('.jpg') or lower_f.endswith('.jpeg')):
            # Si tenemos lista de validacion, verificamos. Si no, usamos todo lo que parezca lamina.
            if valid_images_set is None or f in valid_images_set:
                images.append(os.path.join(folder_path, f))
    
    # Si no hay laminas, no hacer nada
    if not images:
        return

    docx_path = os.path.join(folder_path, "imprimir_lamina_para_sublimar.docx")
    
    # Crear documento
    doc = Document()
    section = doc.sections[0]
    
    # Configurar Pagina Vertical A4 (Estandar Impresora)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    # Margenes para Epson EcoTank L121 (Minimo seguro aprox 3-5mm)
    # Usamos 5mm para aprovechar el maximo de la hoja sin riesgos de corte
    margin_val = Mm(5)
    section.left_margin = margin_val
    section.right_margin = margin_val
    section.top_margin = margin_val
    section.bottom_margin = margin_val
    
    # Detectar tipo de producto por la carpeta
    path_lower = folder_path.lower()
    is_mug_or_cap = 'mug' in path_lower or 'gorra' in path_lower
    
    # Ancho maximo disponible para imprimir
    printable_width = section.page_width - section.left_margin - section.right_margin

    for i, img_path in enumerate(images):
        try:
            if is_mug_or_cap:
                # Medidas especificas solicitadas: 19.5cm x 7.9cm
                doc.add_picture(img_path, width=Cm(19.5), height=Cm(7.9))
            else:
                # Camiseta, Saco u otros: Maximo de la hoja
                doc.add_picture(img_path, width=printable_width)
            
            if i < len(images) - 1:
                doc.add_paragraph() # Salto de linea entre imagenes
        except Exception as e:
            print(f"Error agregando imagen {img_path}: {e}")

    try:
        doc.save(docx_path)
        print(f"Generado documento Word en: {docx_path}")
    except Exception as e:
        print(f"Error guardando Word en {docx_path}: {e}")

def main():
    root = 'img'
    valid_images = load_valid_images()
    
    for dirpath, dirnames, filenames in os.walk(root):
        # Verificar si es una carpeta de producto (contiene laminas)
        if any(f.startswith('lamina_') for f in filenames):
            create_word_for_folder(dirpath, valid_images)

if __name__ == "__main__":
    main()